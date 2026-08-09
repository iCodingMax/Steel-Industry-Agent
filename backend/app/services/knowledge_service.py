"""
文档解析服务（基于LlamaIndex）

本模块提供文档加载、文本切片、知识库管理等功能，是RAG系统的数据准备层。

主要组件：
1. DocumentParserService：文档解析服务
   - 支持多种文档格式加载（PDF、Word、TXT、Markdown）
   - 使用LlamaIndex SentenceSplitter进行文本切片
   - 处理文档并存储切片到数据库

2. KnowledgeBaseService：知识库服务
   - 知识库CRUD操作（创建、查询、更新、删除）
   - 管理知识库配置（名称、描述、切片参数等）

3. DocumentService：文档服务
   - 文档CRUD操作（创建、查询、删除）
   - 管理文档状态（待处理、处理中、已完成、失败）

关键流程：
    文档上传
        │
        ▼
    DocumentService.create()  → 创建文档记录
        │
        ▼
    DocumentParserService.process_document()  → 加载→切片→存储
        │
        ├── load_document()        → 加载文档内容
        ├── split_text()           → 文本切片
        └── 存储到document_segments表
        │
        ▼
    VectorIndexService.build_index()  → 构建向量索引
"""
import os
import json
import asyncio
from typing import List, Optional, Tuple
from pathlib import Path
from loguru import logger

from sqlalchemy import select, delete
from sqlalchemy.ext.asyncio import AsyncSession

from app.models.knowledge import KnowledgeBase, Document, DocumentSegment
from app.schemas.knowledge import (
    KnowledgeBaseCreate,
    KnowledgeBaseUpdate,
    KnowledgeQuery,
    KnowledgeQueryResult,
)
from app.middlewares.exception_handler import BusinessException


class DocumentParserService:
    """
    文档解析服务类

    负责文档的加载、切片和处理，是RAG系统的数据准备核心。
    支持多种文档格式，使用LlamaIndex进行文本切片。

    属性：
        SUPPORTED_FILE_TYPES: 支持的文件类型列表（txt、pdf、md、docx）
    """

    SUPPORTED_FILE_TYPES = ["txt", "pdf", "md", "docx"]

    @staticmethod
    def _get_file_type(file_name: str) -> str:
        """
        获取文件类型

        从文件名中提取扩展名，并验证是否支持该类型。

        :param file_name: 文件名
        :return: 文件类型（小写扩展名）
        :raises BusinessException: 不支持的文件类型时抛出

        验证逻辑：
            1. 提取文件扩展名（去除点号）
            2. 转换为小写
            3. 检查是否在SUPPORTED_FILE_TYPES列表中
            4. 不在列表中则抛出异常
        """
        ext = Path(file_name).suffix.lower().lstrip(".")
        if ext in DocumentParserService.SUPPORTED_FILE_TYPES:
            return ext
        raise BusinessException(code=400, message=f"不支持的文件类型: {ext}")

    @staticmethod
    async def load_document(file_path: str, file_type: str) -> List[str]:
        """
        加载文档内容

        根据文件类型调用相应的加载方法，返回文档的文本段落列表。

        :param file_path: 文件路径
        :param file_type: 文件类型（txt/pdf/md/docx）
        :return: 文档文本段落列表

        支持的文件类型及加载方式：
            - txt: 直接读取全文
            - md: 直接读取全文（保留Markdown格式）
            - pdf: 使用pypdf逐页提取文本
            - docx: 使用python-docx提取段落

        异常处理：
            - 文件不存在或读取失败时抛出BusinessException
            - 记录详细错误日志
        """
        try:
            logger.info(f"开始加载文档: {file_path}, 类型={file_type}")

            if file_type == "txt":
                with open(file_path, "r", encoding="utf-8") as f:
                    content = f.read()
                logger.info(f"TXT文档加载完成: {file_path}, 长度={len(content)}字符")
                return [content]

            elif file_type == "md":
                with open(file_path, "r", encoding="utf-8") as f:
                    content = f.read()
                logger.info(f"Markdown文档加载完成: {file_path}, 长度={len(content)}字符")
                return [content]

            elif file_type == "pdf":
                from pypdf import PdfReader
                reader = PdfReader(file_path)
                pages = []
                for page in reader.pages:
                    text = page.extract_text()
                    if text.strip():
                        pages.append(text)
                logger.info(f"PDF文档加载完成: {file_path}, 共{len(pages)}页")
                return pages

            elif file_type == "docx":
                from docx import Document
                doc = Document(file_path)
                paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
                logger.info(f"DOCX文档加载完成: {file_path}, 共{len(paragraphs)}段落")
                return paragraphs

            else:
                raise BusinessException(code=400, message=f"不支持的文件类型: {file_type}")

        except Exception as e:
            logger.error(f"文档加载失败: {file_path}, 错误: {e}", exc_info=True)
            raise BusinessException(code=500, message=f"文档加载失败: {str(e)}")

    @staticmethod
    def split_text(
        text: str,
        chunk_size: int = 500,
        chunk_overlap: int = 100,
    ) -> List[Tuple[str, int, int]]:
        """
        文本切片

        使用LlamaIndex的SentenceSplitter对文本进行切片，保留切片的位置信息。

        :param text: 原始文本
        :param chunk_size: 切片大小（字符数），默认500
        :param chunk_overlap: 切片重叠长度（字符数），默认100
        :return: [(切片内容, 起始位置, 结束位置)]

        切片逻辑：
            1. 使用SentenceSplitter按句子边界切分文本
            2. 计算每个切片在原始文本中的起始和结束位置
            3. 返回切片内容及位置信息

        注意：
            - chunk_overlap用于保持切片之间的语义连贯性
            - 位置信息用于后续溯源和定位
        """
        from llama_index.core.text_splitter import SentenceSplitter

        splitter = SentenceSplitter(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
        )

        chunks = splitter.split_text(text)
        # 计算每个切片的位置信息
        result = []
        current_pos = 0
        for chunk in chunks:
            start_pos = text.find(chunk[:50], current_pos)  # 查找切片起始位置
            if start_pos == -1:
                start_pos = current_pos
            end_pos = start_pos + len(chunk)
            result.append((chunk, start_pos, end_pos))
            current_pos = end_pos - chunk_overlap  # 考虑重叠

        logger.info(f"文本切片完成: 原文长度{len(text)}, 切片数量{len(result)}, 切片大小{chunk_size}, 重叠{chunk_overlap}")
        return result

    @staticmethod
    async def process_document(
        db: AsyncSession,
        document: Document,
        knowledge_base: KnowledgeBase,
        storage_dir: str = "./storage/documents",
    ) -> int:
        """
        处理文档：加载、切片、存储

        完整的文档处理流程，包括状态更新、内容加载、文本切片和数据库存储。

        :param db: 数据库会话
        :param document: 文档记录对象
        :param knowledge_base: 知识库配置对象（包含切片参数）
        :param storage_dir: 文档存储目录（默认"./storage/documents"）
        :return: 切片数量

        流程步骤：
            1. 更新文档状态为"processing"（处理中）
            2. 加载文档内容（根据文件类型调用相应方法）
            3. 合并所有段落为完整文本
            4. 使用知识库配置的参数进行文本切片
            5. 将切片存储到document_segments表
            6. 更新文档状态为"completed"（已完成）并记录切片数量
            7. 如果处理失败，更新状态为"failed"并记录错误信息

        异常处理：
            - 任何步骤失败时，更新文档状态为"failed"
            - 记录详细错误日志
            - 抛出BusinessException
        """
        try:
            # 步骤1：更新文档状态为处理中
            logger.info(f"开始处理文档: ID={document.id}, 文件名={document.file_name}")
            document.status = "processing"
            await db.commit()

            # 步骤2：加载文档内容
            file_path = document.file_path
            file_type = document.file_type
            paragraphs = await DocumentParserService.load_document(file_path, file_type)

            # 步骤3：合并所有段落，清除PostgreSQL不支持的null字节(\x00)
            full_text = "\n\n".join(paragraphs).replace("\x00", "")
            logger.info(f"文档内容合并完成: 长度={len(full_text)}字符")

            # 步骤4：切片（使用知识库配置的参数）
            chunks = DocumentParserService.split_text(
                full_text,
                chunk_size=knowledge_base.chunk_size,
                chunk_overlap=knowledge_base.chunk_overlap,
            )

            # 步骤5：存储切片到数据库
            for idx, (content, start_char, end_char) in enumerate(chunks):
                segment = DocumentSegment(
                    document_id=document.id,
                    knowledge_base_id=knowledge_base.id,
                    content=content,
                    segment_index=idx,
                    start_char=start_char,
                    end_char=end_char,
                    meta_data=json.dumps({
                        "file_name": document.file_name,
                        "file_type": document.file_type,
                    }),
                )
                db.add(segment)

            # 步骤6：更新文档状态和切片数量
            document.status = "completed"
            document.segment_count = len(chunks)
            await db.commit()

            logger.success(f"文档处理完成: {document.file_name}, 切片数量: {len(chunks)}")
            return len(chunks)

        except Exception as e:
            # 处理失败，更新状态为failed并记录错误信息
            document.status = "failed"
            document.error_message = str(e)
            await db.commit()
            logger.error(f"文档处理失败: {document.file_name}, 错误: {e}", exc_info=True)
            raise BusinessException(code=500, message=f"文档处理失败: {str(e)}")


class KnowledgeBaseService:
    """
    知识库服务类

    提供知识库的CRUD操作，管理知识库配置信息。

    核心功能：
        - 创建知识库（包含名称、描述、切片参数等）
        - 查询知识库（单个查询、列表查询）
        - 更新知识库配置
        - 删除知识库（级联删除关联的文档和切片）
    """

    @staticmethod
    async def create(
        db: AsyncSession,
        data: KnowledgeBaseCreate,
        user_id: Optional[int] = None,
    ) -> KnowledgeBase:
        """
        创建知识库

        创建新的知识库配置，包含名称、描述、嵌入模型和切片参数。

        :param db: 数据库会话
        :param data: 知识库创建参数（Pydantic模型）
        :param user_id: 创建人ID（可选）
        :return: 创建的知识库对象

        默认配置：
            - embedding_model: bge-m3（通过Xinference调用）
            - chunk_size: 500（文本切片大小）
            - chunk_overlap: 100（切片重叠长度）
        """
        kb = KnowledgeBase(
            name=data.name,
            description=data.description,
            embedding_model=data.embeddingModel,
            chunk_size=data.chunkSize,
            chunk_overlap=data.chunkOverlap,
            created_by=user_id,
        )
        db.add(kb)
        await db.commit()
        await db.refresh(kb)
        logger.info(f"创建知识库: {kb.name} (ID: {kb.id})")
        return kb

    @staticmethod
    async def get_by_id(db: AsyncSession, kb_id: int) -> Optional[KnowledgeBase]:
        """
        根据ID获取知识库

        :param db: 数据库会话
        :param kb_id: 知识库ID
        :return: 知识库对象或None
        """
        stmt = select(KnowledgeBase).where(KnowledgeBase.id == kb_id)
        result = await db.execute(stmt)
        return result.scalar_one_or_none()

    @staticmethod
    async def get_all(db: AsyncSession, skip: int = 0, limit: int = 100) -> List[KnowledgeBase]:
        """
        获取所有知识库

        :param db: 数据库会话
        :param skip: 跳过数量（默认0）
        :param limit: 返回数量（默认100）
        :return: 知识库列表
        """
        stmt = select(KnowledgeBase).offset(skip).limit(limit)
        result = await db.execute(stmt)
        return list(result.scalars().all())

    @staticmethod
    async def update(
        db: AsyncSession,
        kb_id: int,
        data: KnowledgeBaseUpdate,
    ) -> Optional[KnowledgeBase]:
        """
        更新知识库

        更新知识库配置，支持部分字段更新。

        :param db: 数据库会话
        :param kb_id: 知识库ID
        :param data: 更新参数（Pydantic模型）
        :return: 更新后的知识库对象或None

        字段映射：
            - embeddingModel → embedding_model
            - chunkSize → chunk_size
            - chunkOverlap → chunk_overlap
        """
        kb = await KnowledgeBaseService.get_by_id(db, kb_id)
        if not kb:
            raise BusinessException(code=404, message="知识库不存在")

        update_data = data.model_dump(exclude_unset=True)
        # 处理字段名映射（驼峰命名→蛇形命名）
        field_mapping = {
            "embeddingModel": "embedding_model",
            "chunkSize": "chunk_size",
            "chunkOverlap": "chunk_overlap",
        }
        for key, value in update_data.items():
            db_key = field_mapping.get(key, key)
            if hasattr(kb, db_key):
                setattr(kb, db_key, value)

        await db.commit()
        await db.refresh(kb)
        logger.info(f"更新知识库: {kb.name} (ID: {kb.id})")
        return kb

    @staticmethod
    async def delete(db: AsyncSession, kb_id: int) -> None:
        """
        删除知识库

        删除知识库及其关联的所有文档和切片。

        :param db: 数据库会话
        :param kb_id: 知识库ID
        :raises BusinessException: 知识库不存在时抛出

        删除顺序：
            1. 删除关联的文档切片（document_segments）
            2. 删除关联的文档（documents）
            3. 删除知识库（knowledge_bases）
        """
        kb = await KnowledgeBaseService.get_by_id(db, kb_id)
        if not kb:
            raise BusinessException(code=404, message="知识库不存在")

        # 删除关联的切片
        await db.execute(
            delete(DocumentSegment).where(DocumentSegment.knowledge_base_id == kb_id)
        )
        # 删除关联的文档
        await db.execute(
            delete(Document).where(Document.knowledge_base_id == kb_id)
        )
        # 删除知识库
        await db.delete(kb)
        await db.commit()
        logger.info(f"删除知识库: {kb.name} (ID: {kb_id})")


class DocumentService:
    """
    文档服务类

    提供文档的CRUD操作，管理文档的上传、状态和删除。

    核心功能：
        - 创建文档记录（文件名、路径、类型等）
        - 查询文档（单个查询、按知识库查询）
        - 删除文档（级联删除关联的切片和文件）
    """

    @staticmethod
    async def create(
        db: AsyncSession,
        knowledge_base_id: int,
        file_name: str,
        file_path: str,
        file_type: str,
        file_size: Optional[int] = None,
    ) -> Document:
        """
        创建文档记录

        在数据库中创建新的文档记录，初始状态为"pending"（待处理）。

        :param db: 数据库会话
        :param knowledge_base_id: 所属知识库ID
        :param file_name: 文件名
        :param file_path: 文件存储路径
        :param file_type: 文件类型
        :param file_size: 文件大小（字节，可选）
        :return: 创建的文档对象

        初始状态：
            - status: pending（待处理）
            - segment_count: 0（切片数量）
        """
        doc = Document(
            knowledge_base_id=knowledge_base_id,
            file_name=file_name,
            file_path=file_path,
            file_type=file_type,
            file_size=file_size,
            status="pending",
        )
        db.add(doc)
        await db.commit()
        await db.refresh(doc)
        logger.info(f"创建文档记录: {file_name} (ID: {doc.id})")
        return doc

    @staticmethod
    async def get_by_id(db: AsyncSession, doc_id: int) -> Optional[Document]:
        """
        根据ID获取文档

        :param db: 数据库会话
        :param doc_id: 文档ID
        :return: 文档对象或None
        """
        stmt = select(Document).where(Document.id == doc_id)
        result = await db.execute(stmt)
        return result.scalar_one_or_none()

    @staticmethod
    async def get_by_knowledge_base(
        db: AsyncSession,
        kb_id: int,
        skip: int = 0,
        limit: int = 100,
    ) -> List[Document]:
        """
        获取知识库下的所有文档

        :param db: 数据库会话
        :param kb_id: 知识库ID
        :param skip: 跳过数量（默认0）
        :param limit: 返回数量（默认100）
        :return: 文档列表
        """
        stmt = select(Document).where(Document.knowledge_base_id == kb_id).offset(skip).limit(limit)
        result = await db.execute(stmt)
        return list(result.scalars().all())

    @staticmethod
    async def delete(db: AsyncSession, doc_id: int) -> None:
        """
        删除文档

        删除文档记录、关联的切片和物理文件。

        :param db: 数据库会话
        :param doc_id: 文档ID
        :raises BusinessException: 文档不存在时抛出

        删除顺序：
            1. 删除关联的文档切片（document_segments）
            2. 删除物理文件（如果存在）
            3. 删除文档记录（documents）
        """
        doc = await DocumentService.get_by_id(db, doc_id)
        if not doc:
            raise BusinessException(code=404, message="文档不存在")

        # 删除关联的切片
        await db.execute(
            delete(DocumentSegment).where(DocumentSegment.document_id == doc_id)
        )
        # 删除物理文件
        if os.path.exists(doc.file_path):
            os.remove(doc.file_path)
            logger.info(f"删除物理文件: {doc.file_path}")
        # 删除文档记录
        await db.delete(doc)
        await db.commit()
        logger.info(f"删除文档: {doc.file_name} (ID: {doc_id})")


# 服务实例（供其他模块调用）
document_parser_service = DocumentParserService()
knowledge_base_service = KnowledgeBaseService()
document_service = DocumentService()